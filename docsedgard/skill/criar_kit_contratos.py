from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = r"D:\AI Jedgard\outputs\produtos_ebook\kit_contratos_sem_surpresa.docx"

doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.75)
section.bottom_margin = Inches(0.75)
section.left_margin = Inches(0.8)
section.right_margin = Inches(0.8)

styles = doc.styles
styles["Normal"].font.name = "Arial"
styles["Normal"].font.size = Pt(10)
styles["Normal"].paragraph_format.space_after = Pt(6)
styles["Normal"].paragraph_format.line_spacing = 1.08

for name, size, color in [
    ("Heading 1", 17, RGBColor(6, 9, 15)),
    ("Heading 2", 13, RGBColor(153, 27, 27)),
    ("Heading 3", 11, RGBColor(6, 9, 15)),
]:
    st = styles[name]
    st.font.name = "Arial"
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = color
    st.paragraph_format.space_before = Pt(10)
    st.paragraph_format.space_after = Pt(5)

def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)

def border_table(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "6")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), "CBD5E1")
        borders.append(tag)
    tbl_pr.append(borders)

def p(text="", style=None, bold=False, color=None, align=None):
    para = doc.add_paragraph(style=style)
    run = para.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    if color:
        run.font.color.rgb = color
    if align:
        para.alignment = align
    return para

def bullet(text):
    para = doc.add_paragraph(style=None)
    para.style = styles["List Bullet"]
    para.paragraph_format.left_indent = Inches(0.25)
    para.paragraph_format.first_line_indent = Inches(-0.12)
    run = para.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(10)

def clause(title, body):
    p(title, style="Heading 3")
    p(body)

def two_col_table(rows):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(2.2)
    table.columns[1].width = Inches(4.7)
    hdr = table.rows[0].cells
    hdr[0].text = "Campo"
    hdr[1].text = "Preencher"
    for c in hdr:
        shade(c, "06090F")
        c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for par in c.paragraphs:
            for r in par.runs:
                r.font.color.rgb = RGBColor(255, 255, 255)
                r.font.bold = True
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
        shade(cells[0], "F8FAFC")
        for c in cells:
            c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    border_table(table)
    return table

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("KIT CONTRATOS SEM SURPRESA")
r.bold = True
r.font.name = "Arial"
r.font.size = Pt(24)
r.font.color.rgb = RGBColor(6, 9, 15)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = subtitle.add_run("Modelos, campos e clausulas de atencao para contratar obra com mais controle")
r.font.name = "Arial"
r.font.size = Pt(11)
r.font.color.rgb = RGBColor(71, 85, 105)

p("Aviso importante", style="Heading 1")
p("Este material e um guia educativo e operacional. Ele nao substitui advogado, engenheiro, arquiteto ou responsavel tecnico. Antes de assinar qualquer contrato real, valide o texto com profissionais habilitados e adapte as clausulas a legislacao local, ao escopo e ao risco da obra.", bold=True, color=RGBColor(153, 27, 27))

p("Como usar este kit", style="Heading 1")
for item in [
    "Use os campos como roteiro para montar ou revisar propostas de obra.",
    "Nunca aceite escopo generico como 'mao de obra completa' sem lista de entregas.",
    "Guarde anexos: projetos, memoriais, cronograma, orcamento, fotos e mensagens.",
    "Toda alteracao de preco ou prazo deve virar aditivo antes da execucao.",
]:
    bullet(item)

p("Modelo 1 - Contratacao de servico de obra/reforma", style="Heading 1")
two_col_table([
    ("Contratante", "[Nome, CPF/CNPJ, endereco]"),
    ("Contratado", "[Nome, CPF/CNPJ, CREA/CAU se aplicavel]"),
    ("Endereco da obra", "[Endereco completo]"),
    ("Objeto", "[Descricao objetiva do servico]"),
    ("Valor total", "R$ [valor]"),
    ("Prazo", "[data de inicio] a [data de termino]"),
    ("Responsavel tecnico", "[nome, registro, ART/RRT]"),
])

clause("1. Objeto e escopo", "O contratado executara exclusivamente os servicos descritos no Anexo I - Escopo Tecnico. Qualquer item nao listado sera considerado fora do escopo e dependera de aprovacao previa por escrito.")
clause("2. Documentos integrantes", "Fazem parte deste contrato: proposta comercial, escopo tecnico, cronograma fisico-financeiro, memorial descritivo, projetos, ART/RRT quando aplicavel e registros de alteracoes aprovadas.")
clause("3. Prazo e condicoes de inicio", "O prazo somente comecara a contar apos liberacao do local, entrega dos projetos necessarios, pagamento inicial previsto e disponibilizacao das condicoes minimas de execucao.")
clause("4. Pagamento por medicao", "Os pagamentos serao vinculados a etapas executadas e conferidas. O contratante podera reter valores proporcionais a servicos incompletos, divergentes ou sem comprovacao.")
clause("5. Materiais", "Definir se materiais serao fornecidos pelo contratante ou contratado. Marcas, especificacoes, quantidades e substituicoes deverao ser aprovadas por escrito.")
clause("6. Alteracoes de escopo", "Nenhuma alteracao sera executada sem aditivo contendo descricao, impacto de preco, impacto de prazo e aceite das partes.")
clause("7. Qualidade e correcao", "Servicos executados em desacordo com projeto, norma, memorial ou boa tecnica deverao ser corrigidos pelo contratado sem custo adicional, quando comprovada responsabilidade.")
clause("8. Entrega tecnica", "A entrega final devera incluir limpeza basica, conferencia do escopo, relacao de pendencias, garantias, manuais e documentos tecnicos aplicaveis.")

doc.add_page_break()
p("Modelo 2 - Aditivo de alteracao de escopo", style="Heading 1")
two_col_table([
    ("Contrato original", "[Numero/data]"),
    ("Solicitante", "[Contratante ou contratado]"),
    ("Alteracao", "[Descricao clara do que muda]"),
    ("Motivo", "[Necessidade tecnica, escolha do cliente, imprevisto etc.]"),
    ("Acrescimo/desconto", "R$ [valor]"),
    ("Impacto no prazo", "[dias a mais/a menos]"),
    ("Aprovacao", "[assinaturas/data]"),
])
clause("Clausula de aceite do aditivo", "As partes reconhecem que a alteracao descrita neste aditivo modifica o escopo original. A execucao somente devera iniciar apos assinatura ou aceite formal registrado por escrito.")

p("Modelo 3 - Termo de entrega e pendencias", style="Heading 1")
two_col_table([
    ("Obra/servico", "[Descricao]"),
    ("Data de vistoria", "[Data]"),
    ("Participantes", "[Nomes]"),
    ("Status geral", "[Aprovado / Aprovado com pendencias / Reprovado]"),
    ("Prazo para pendencias", "[Data limite]"),
    ("Retencao financeira", "R$ [se aplicavel]"),
])

items = [
    "Conferir escopo contratado item por item.",
    "Registrar fotos das areas vistoriadas.",
    "Listar pendencias com prazo e responsavel.",
    "Conferir limpeza, remocao de entulho e protecao de acabamentos.",
    "Receber notas fiscais, garantias, manuais e documentos tecnicos.",
]
p("Checklist de entrega", style="Heading 2")
for item in items:
    bullet(item)

p("Clausulas que merecem atencao maxima", style="Heading 1")
table = doc.add_table(rows=1, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = "Tema"
hdr[1].text = "Risco"
hdr[2].text = "Como blindar"
for c in hdr:
    shade(c, "06090F")
    for par in c.paragraphs:
        for run in par.runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.bold = True
rows = [
    ("Escopo", "Contratado diz que item nao estava incluso.", "Anexar lista detalhada de servicos, materiais e exclusoes."),
    ("Prazo", "Atraso sem consequencia clara.", "Definir marco inicial, causas justificadas e multa/retencao."),
    ("Pagamento", "Cliente paga antes da entrega real.", "Vincular pagamento a medicao e aceite da etapa."),
    ("Materiais", "Troca por material inferior.", "Exigir marcas, especificacoes ou aprovacao previa."),
    ("Garantia", "Ninguem assume defeito depois.", "Definir garantia, prazo de resposta e forma de acionamento."),
    ("Responsabilidade tecnica", "Obra sem responsavel habilitado.", "Exigir ART/RRT quando aplicavel."),
]
for row in rows:
    cells = table.add_row().cells
    for i, value in enumerate(row):
        cells[i].text = value
border_table(table)

p("Perguntas antes de assinar", style="Heading 1")
for item in [
    "O que exatamente esta incluso e o que esta excluido?",
    "Quem compra os materiais e quem responde por perdas?",
    "Como sera feita a medicao de cada etapa?",
    "O que acontece se houver atraso?",
    "Como serao aprovadas mudancas de escopo?",
    "Existe responsavel tecnico e documento de responsabilidade?",
    "Qual parte do pagamento fica retida ate a correcao de pendencias?",
]:
    bullet(item)

doc.save(OUT)
