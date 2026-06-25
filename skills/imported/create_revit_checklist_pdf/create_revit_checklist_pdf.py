from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_DOCX = r"D:\AI Jedgard\Checklist_Revit_Template_Edgard.docx"
OUT_PDF = r"D:\AI Jedgard\Checklist_Revit_Template_Edgard.pdf"


sections = [
    ("1. Base do Template", [
        "principal.rte como template principal.",
        "Base parecida com Castanhari/Jedi.",
        "Usar filtros como base de organizacao visual.",
        "Padrao proprio Edgard.",
        "Organizar por etapa da obra.",
        "Deixar leve, mas fazer limpeza pesada somente no final.",
        "Nao excluir nada sem listar e confirmar.",
    ]),
    ("2. Pranchas", [
        "A01 - Implantacao / Situacao.",
        "A02 - Plantas Baixas.",
        "A03 - Plantas Humanizadas.",
        "A04 - Fachadas.",
        "A05 - Areas.",
        "A06 - Estrutural / Hidraulica / Eletrica / 3D.",
        "Inserir vistas nos lugares corretos.",
        "Colocar carimbo Edgard no final.",
        "Remover carimbos extras somente no final e com confirmacao.",
    ]),
    ("3. Vistas", [
        "Planta baixa terreo.",
        "Planta baixa 2 pavimento.",
        "Terreo planta humanizada.",
        "2 pavimento planta humanizada.",
        "Situacao.",
        "4 fachadas: frontal, posterior, esquerda e direita.",
        "Calculo de areas terreo.",
        "Calculo de areas 2 pavimento.",
        "3D estrutural.",
        "Plantas estruturais.",
        "Plantas hidraulicas.",
        "Plantas eletricas.",
        "Renomear vistas sem nome/interrogacao.",
        "Retirar acentos problemáticos quando necessario.",
        "Remover nomes Template Thiago Castanhari.",
    ]),
    ("4. Templates de Vista", [
        "Arquitetura.",
        "Estrutura.",
        "Eletrica.",
        "Hidraulica.",
        "Planta baixa padrao.",
        "Planta humanizada em melhor qualidade.",
        "Anotacoes principais.",
        "Cotas.",
        "Ambientes.",
        "Mobiliario.",
        "Template combinado: arquitetura + estrutura + eletrica + hidraulica.",
        "Aplicar templates nas vistas principais.",
    ]),
    ("5. Filtros", [
        "Aproveitar filtros do Castanhari/Jedi.",
        "Adaptar filtros para padrao Edgard.",
        "Configurar filtros de arquitetura.",
        "Configurar filtros de estrutura.",
        "Configurar filtros de eletrica.",
        "Configurar filtros de hidraulica.",
        "Configurar filtros de paredes.",
        "Configurar filtros de lajes.",
        "Configurar filtros de mobiliario.",
        "Configurar filtros de elementos novos/remover/deslocar.",
        "Configurar filtros de niveis auxiliares.",
        "Configurar filtros de cortes/fases.",
        "Verificar se os filtros estao aplicados nos templates.",
    ]),
    ("6. Legendas", [
        "Criar legendas principais.",
        "Legenda de simbolos e anotacoes.",
        "Legenda de paredes/acabamentos.",
        "Legenda de esquadrias.",
        "Legenda eletrica.",
        "Legenda hidraulica.",
        "Legenda de materiais.",
        "Preencher legendas com conteudo real.",
    ]),
    ("7. Tabelas", [
        "Tabela de areas por pavimento.",
        "Trazer informacoes das plantas de calculo de area terreo e 2 pavimento.",
        "Incluir outros pavimentos se existirem.",
        "Tabela de quantitativo geral.",
        "Tabela de materiais brutos.",
        "Tabela de paredes.",
        "Tabela de telhados.",
        "Tabela de portas.",
        "Tabela de janelas.",
        "Tabela de pisos.",
        "Tabela de pilares.",
        "Tabela de vigas.",
        "Tabela de laje.",
        "Tabela de madeiramento.",
        "Tabela de steel frame/estrutura metalica quando houver.",
    ]),
    ("8. Insumos e Calculos", [
        "Ferragem de alicerce.",
        "Ferragem de colunas.",
        "Ferragem de vigas.",
        "Concreto para alicerce.",
        "Concreto para vigas baldrame.",
        "Cimento por saco.",
        "Areia por m3.",
        "Areia por caminhao de 12 m3.",
        "Pedra por m3.",
        "Pedra por caminhao.",
        "Plastificante/liquido de liga por balde de 18 litros.",
        "Reboco: cimento, areia e plastificante.",
        "Bloco 6 furos 19 x 11,5 x 24.",
        "Outros tipos de bloco.",
        "Massa corrida por lata.",
        "Tinta por lata.",
        "Laje em m3.",
        "Madeiramento conforme telhado.",
        "Telhas.",
        "Estrutura metalica.",
        "Steel frame.",
        "Perguntar antes de finalizar se falta algum item.",
    ]),
    ("9. Familias", [
        "Juntar familias uteis dos templates no principal.",
        "Verificar familias duplicadas.",
        "Organizar pasta de familias.",
        "Verificar familias com nome mas sem imagem/preview.",
        "Tentar achar imagem/preview correta quando fizer sentido.",
        "Marcar familias ruins para exclusao.",
        "Excluir familias menos uteis somente no final.",
        "Fazer lista antes de excluir.",
        "Manter backup.",
    ]),
    ("10. Materiais", [
        "Listar materiais.",
        "Verificar materiais sem imagem/textura.",
        "Associar imagem correta quando fizer sentido.",
        "Marcar materiais ruins/inuteis para revisao.",
        "Preparar materiais para planta humanizada.",
        "Preparar materiais para quantitativo.",
    ]),
    ("11. Paredes e Telhados", [
        "Configurar tipos de parede para quantitativo.",
        "Padronizar paredes por uso.",
        "Configurar tipos de telhado para quantitativo.",
        "Telhados com madeiramento.",
        "Telhados com ferragem/estrutura metalica.",
        "Telhados com steel frame quando aplicavel.",
    ]),
    ("12. Automacoes PyRevit / Botoes", [
        "Preparar Projeto.",
        "Gerar Pranchas.",
        "Gerar Planta Humanizada.",
        "Gerar Cortes e Fachadas.",
        "Gerar 3D.",
        "Gerar Percurso Animado.",
        "Gerar Estudo Solar.",
        "Gerar Quantitativos.",
        "Exportar Pacote.",
        "Gerar Memorial e Contrato.",
    ]),
    ("13. Planta Humanizada", [
        "Ao desenhar planta terreo, humanizada ja ficar bonita.",
        "Melhor qualidade visual sem render pesado.",
        "Configurar materiais, sombras e estilo grafico.",
        "Exportar imagem para IA/render externo se necessario.",
    ]),
    ("14. Norte e Estudo Solar", [
        "Configurar norte verdadeiro.",
        "Configurar norte do projeto quando necessario.",
        "Criar estudo solar.",
        "Gerar vistas solares padrao.",
        "Gerar imagens de estudo solar.",
    ]),
    ("15. Cortes, Fachadas, 3D e Percursos", [
        "Gerar cortes padroes de arquitetura.",
        "Gerar fachadas frontal/posterior/esquerda/direita.",
        "Criar 3D padrao.",
        "Criar 3D estrutural.",
        "Criar percurso animado automaticamente.",
        "Criar vistas por ambiente se possivel.",
    ]),
    ("16. Integracao com Plataforma", [
        "Exportar produto pronto do Revit.",
        "Exportar PDF.",
        "Exportar imagens.",
        "Exportar tabelas.",
        "Exportar modelo/BIM.",
        "Enviar para plataforma automaticamente.",
        "Plataforma gerar projeto, contrato, documentacao, memoriais, render IA, analise BIM e quantitativos.",
    ]),
    ("17. Ordem Atual Combinada", [
        "Fazer primeiro estrutura do template.",
        "Depois pranchas, vistas, templates, filtros, legendas e tabelas.",
        "Depois automacoes.",
        "Deixar para o final: carimbo, exclusao de familias, limpeza pesada, duplicidades e materiais ruins.",
    ]),
    ("Ja Feito", [
        "Pranchas A01-A06 criadas.",
        "14 vistas principais posicionadas.",
        "Templates EDGARD principais existem.",
        "Templates complementares criados.",
        "Templates aplicados nas vistas principais.",
        "6 legendas base criadas.",
        "Tabelas EDGARD iniciais criadas.",
        "Memoria atualizada.",
    ]),
    ("Ainda Falta Bastante", [
        "Filtros configurados de verdade.",
        "Conteudo das legendas.",
        "Tabelas com formulas de insumos.",
        "Paredes/telhados parametrizados.",
        "Planta humanizada visualmente ajustada.",
        "Cortes automaticos.",
        "Percurso animado.",
        "Estudo solar.",
        "Automacoes pyRevit.",
        "Carimbo Edgard definitivo.",
        "Limpeza final com aprovacao.",
    ]),
]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, size=9.0, color="111111"):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)


doc = Document()
section = doc.sections[0]
section.orientation = WD_ORIENT.PORTRAIT
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.55)
section.bottom_margin = Inches(0.55)
section.left_margin = Inches(0.55)
section.right_margin = Inches(0.55)

styles = doc.styles
styles["Normal"].font.name = "Arial"
styles["Normal"].font.size = Pt(9)
styles["Normal"].paragraph_format.space_after = Pt(3)
styles["Normal"].paragraph_format.line_spacing = 1.05

for style_name, size, color in [
    ("Heading 1", 12, "0B2545"),
    ("Heading 2", 10.5, "1F4D78"),
]:
    style = styles[style_name]
    style.font.name = "Arial"
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(8)
    style.paragraph_format.space_after = Pt(3)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(2)
r = title.add_run("CHECKLIST DO TEMPLATE REVIT EDGARD")
r.bold = True
r.font.name = "Arial"
r.font.size = Pt(16)
r.font.color.rgb = RGBColor.from_string("0B2545")

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.paragraph_format.space_after = Pt(8)
r = subtitle.add_run("Sequencia de trabalho para o principal.rte | Versao para impressao")
r.font.name = "Arial"
r.font.size = Pt(9)
r.font.color.rgb = RGBColor.from_string("555555")

note = doc.add_paragraph()
note.paragraph_format.space_after = Pt(8)
run = note.add_run("Regra de seguranca: exclusao de familias, limpeza pesada e carimbo definitivo ficam para o final, sempre com lista e confirmacao antes.")
run.bold = True
run.font.name = "Arial"
run.font.size = Pt(8.5)
run.font.color.rgb = RGBColor.from_string("7A5A00")

for heading, items in sections:
    doc.add_heading(heading, level=1)
    table = doc.add_table(rows=1, cols=3)
    table.autofit = False
    table.allow_autofit = False
    widths = [Inches(0.35), Inches(6.45), Inches(0.55)]
    hdr = table.rows[0].cells
    for idx, label in enumerate(["OK", "Item", "Resp."]):
        hdr[idx].width = widths[idx]
        set_cell_shading(hdr[idx], "E8EEF5")
        set_cell_text(hdr[idx], label, bold=True, size=8.5, color="0B2545")
    for item in items:
        cells = table.add_row().cells
        for i, w in enumerate(widths):
            cells[i].width = w
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_text(cells[0], "[  ]", size=8.6, color="333333")
        set_cell_text(cells[1], item, size=8.6)
        set_cell_text(cells[2], "", size=8.6)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run("Checklist Revit Edgard | documento de controle interno")
run.font.name = "Arial"
run.font.size = Pt(8)
run.font.color.rgb = RGBColor.from_string("666666")

doc.save(OUT_DOCX)


def build_pdf():
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import (
        BaseDocTemplate,
        Frame,
        KeepTogether,
        PageTemplate,
        Paragraph,
        Spacer,
        Table,
        TableStyle,
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "TitleEdgard",
        parent=styles["Title"],
        fontName="Helvetica-Bold",
        fontSize=16,
        leading=19,
        alignment=TA_CENTER,
        textColor=colors.HexColor("#0B2545"),
        spaceAfter=3,
    )
    subtitle_style = ParagraphStyle(
        "SubtitleEdgard",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=8.5,
        leading=10,
        alignment=TA_CENTER,
        textColor=colors.HexColor("#555555"),
        spaceAfter=8,
    )
    note_style = ParagraphStyle(
        "NoteEdgard",
        parent=styles["Normal"],
        fontName="Helvetica-Bold",
        fontSize=8.5,
        leading=10.5,
        textColor=colors.HexColor("#7A5A00"),
        spaceAfter=8,
    )
    h_style = ParagraphStyle(
        "HeadingEdgard",
        parent=styles["Heading2"],
        fontName="Helvetica-Bold",
        fontSize=10.4,
        leading=12,
        textColor=colors.HexColor("#0B2545"),
        spaceBefore=6,
        spaceAfter=3,
    )
    item_style = ParagraphStyle(
        "ItemEdgard",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=8.0,
        leading=9.3,
        alignment=TA_LEFT,
        textColor=colors.HexColor("#111111"),
    )
    small_style = ParagraphStyle(
        "SmallEdgard",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=7.5,
        leading=9,
        alignment=TA_CENTER,
        textColor=colors.HexColor("#111111"),
    )

    def footer(canvas, doc_obj):
        canvas.saveState()
        canvas.setFont("Helvetica", 7)
        canvas.setFillColor(colors.HexColor("#666666"))
        canvas.drawCentredString(
            4.25 * inch,
            0.32 * inch,
            "Checklist Revit Edgard | controle interno | pagina %d" % doc_obj.page,
        )
        canvas.restoreState()

    pdf = BaseDocTemplate(
        OUT_PDF,
        pagesize=letter,
        leftMargin=0.55 * inch,
        rightMargin=0.55 * inch,
        topMargin=0.48 * inch,
        bottomMargin=0.48 * inch,
    )
    frame = Frame(pdf.leftMargin, pdf.bottomMargin, pdf.width, pdf.height, id="normal")
    pdf.addPageTemplates([PageTemplate(id="edgard", frames=[frame], onPage=footer)])

    story = [
        Paragraph("CHECKLIST DO TEMPLATE REVIT EDGARD", title_style),
        Paragraph("Sequencia de trabalho para o principal.rte | Versao para impressao", subtitle_style),
        Paragraph(
            "Regra de seguranca: exclusao de familias, limpeza pesada e carimbo definitivo ficam para o final, sempre com lista e confirmacao antes.",
            note_style,
        ),
    ]

    col_widths = [0.32 * inch, 6.42 * inch, 0.55 * inch]
    for heading, items in sections:
        rows = [
            [
                Paragraph("OK", small_style),
                Paragraph("Item", small_style),
                Paragraph("Resp.", small_style),
            ]
        ]
        for item in items:
            rows.append([
                Paragraph("[  ]", small_style),
                Paragraph(item, item_style),
                Paragraph("", small_style),
            ])
        table = Table(rows, colWidths=col_widths, repeatRows=1, hAlign="LEFT")
        table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E8EEF5")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#0B2545")),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#C9D3DF")),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("ALIGN", (0, 0), (0, -1), "CENTER"),
            ("ALIGN", (2, 0), (2, -1), "CENTER"),
            ("LEFTPADDING", (0, 0), (-1, -1), 4),
            ("RIGHTPADDING", (0, 0), (-1, -1), 4),
            ("TOPPADDING", (0, 0), (-1, -1), 3),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
        ]))
        story.append(KeepTogether([Paragraph(heading, h_style), table]))
        story.append(Spacer(1, 5))

    pdf.build(story)


build_pdf()
print(OUT_DOCX)
print(OUT_PDF)
